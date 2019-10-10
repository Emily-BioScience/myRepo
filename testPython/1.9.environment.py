# -*- coding UTF-8 -*-

# 数据处理
import numpy as np
import pandas as pd
import scipy as sp

# 数据可视化
import matplotlib.pyplot as plt
plt.switch_backend('agg')
# import seaborn
# import mayavi

# 文本处理
from PyPDF2 import PdfFileReader, PdfFileMerger
from docx import Document
import nltk

# 机器学习
from sklearn import datasets
import tensorflow as tf
import mxnet

# 页面爬取
import requests
import scrapy
import pyspider

# 页面提取
import beautifulsoup4 as bs4
import re  # 正则表达式
# import python-goose  # 文章/视频元数据

# 网站开发
import Django  # 最流行的web应用框架，MTV模式，模型（model），模板（template），视图（views）
import pyramid  # 简单方便构建web系统的应用框架，起步简单，可扩展性很好
import flask    # 简单，规模小，快速，Django > pyramid > Flask

# 有意思的
import werobot  # 微信公众号开发框架，微信机器人
import aip  # 百度AI开放平台接口，语音、人脸、OCR、NLP、知识图谱、图像搜索
import myqr  # MyQR，生成基本二维码、艺术二维码、动态二维码

# 图形用户界面
import PyQt5  # 创建Qt5程序的Python API接口，工业级使用
import PyGObject  # 跨平台用户图形界面GUI框架

# 游戏开发
import PyGame  # 基于SDL的简单游戏开发功能及实现引擎，入门学习库
import Panda3D  # 开源、跨平台的3D渲染和游戏开发库，迪士尼开发
import cocos2d  # 构建2D游戏和图形界面交互式应用的框架

# 虚拟现实
import VRZero  # 针对树莓派开发VR应用，嵌入式硬件
import pyovr  # 针对Oculus VR设备，Python+设备
import vizrd  # 基于python的通用VR开发引擎

# 空间艺术





def usenumpy():
    a = np.array([0, 1, 2, 3, 4])
    b = np.array([9, 8, 7, 6, 5])
    c = a**2 + b**3  # 将数组看成一个变量，n维数组可以看做对象，减少for循环
    print(">>>>>> numpy\n", c)
    # python接口使用，C语言实现，计算速度优异
    # python数据分析及科学计算的基础库，支撑pandas等
    # 提供直接的矩阵运算、广播函数、线性代数等功能


def usepandas():
    pass
    # 提供数据结构和数据分析工具
    # 理解数据类型与索引的关系，操作索引即操作数据，基于numpy库进行开发
    # python最主要的数据分析功能库，基于numpy进行开发
    # Series = 索引 + 一维数据
    # DataFrame = 行列索引 + 二维数据


def usescipy():
    pass
    # 数学、科学和工程计算功能库
    # 提供了一批数学算法及工程数据运算功能
    # 类似Matlab，可用于如傅里叶变换、信号处理等应用
    # python最主要的科学计算功能库，基于numpy开发
    # 傅里叶变换类，信号处理类，线性代数类，图像处理类，稀疏图运算类，稀疏图压缩类


def usematplotlib(outfile):

    digits = datasets.load_digits()
    fig = plt.figure(figsize=(6, 6))
    fig.subplots_adjust(left=0, right=1, bottom=0, top=1, hspace=0.05, wspace=0.05)
    for i in range(64):
        ax = fig.add_subplot(8, 8, i+1, xticks=[], yticks=[])
        ax.imshow(digits.images[i], cmap=plt.cm.binary, interpolation='nearest')
        ax.text(0, 7, str(digits.target[i]))
    # plt.show()
    plt.savefig(outfile)


def useseaborn():
    pass


def usepythondocx(outfile):
    doc = Document()
    doc.add_heading('test', 0)
    records = ((3, 'test', 'a'), (4, 'haha', 'b'), (5, 'list', 'c'))
    table = doc.add_table(rows=1, cols=3)
    heading = table.rows[0].cells
    heading[0].text = 'id'
    heading[1].text = 'name'
    heading[2].text = 'type'
    for id, name, type in records:
        cells = table.add_row().cells
        cells[0].text = str(id)
        cells[1].text = name
        cells[2].text = type
    doc.save(outfile)


def usepypdf2(infile1, infile2, outfile):
    merger = PdfFileMerger()
    input1 = open(infile1, 'rb')
    input2 = open(infile2, 'rb')
    merger.append(fileobj = input1, pages = (0, 2))
    merger.merge(position = 2, fileobj = input2, pages = (0, 1))
    output = open(outfile, 'wb')
    merger.write(output)


def usenltk():
    sentence = """At eight o'clock on Thursday morning
    ... Arthur didn't feel very good."""
    # tokens = nltk.word_tokenize(sentence)
    # print(tokens)


def usesklearn():
    digits = datasets.load_digits()
    print(">>>>>>>> keys  \n", digits.keys())
    print(">>>>>>>> data  \n", digits.data.shape)
    print(">>>>>>>> target  \n", len(digits.target))
    print(">>>>>>>> DESCR  \n", len(digits.DESCR))


def radarChart():
    pass


def userequests(url, outfile):
    r = requests.get(url, auth=('user', 'pass'))
    r.status_code
    r.headers['content-type']
    r.encoding
    fo = open(outfile, 'w')
    fo.write(r.text)



if __name__ == '__main__':
    usenumpy()
    usepandas()
    usescipy()
    usematplotlib('output/1.9.demo.jpg')
    useseaborn()
    usepythondocx('output/1.9.demo.docx')
    usepypdf2('data/1.9.file1.pdf', 'data/1.9.file2.pdf', 'output/1.9.demo.pdf')
    usenltk()
    usesklearn()
    radarChart()
    userequests('https://www.nature.com/articles/s41467-019-09186-x', 'output/1.9.demo.txt')


# 数据表示，合适方式用程序表达数据
# 数据清洗，数据归一化，数据转换，异常值处理
# 数据统计，数据的概要理解，数量、分布、中位数等
# 数据可视化，直观展示数据内涵的方式
# 数据挖掘，从数据分析获得知识、产生数据外的价值
# 人工智能，数据/语言/图像/视觉等方面深度分析与决策
# 目标（目标感，寻觅之） + 沉浸（沉浸感，思考之） + 熟练（反复练习，熟练之）


